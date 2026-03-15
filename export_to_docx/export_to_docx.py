"""
title: Export to DOCX
description: Export Assistant Message to DOCX
version: 1.0.0
author: Helmi Chaouachi
git_url: https://github.com/Helmi97/open-webui-extensions/tree/main/export_to_docx
icon_url: https://www.svgrepo.com/show/452072/ms-word.svg
required_open_webui_version: 0.8.0
requirements: docxtpl,markdown,requests,beautifulsoup4,python-docx
"""

from __future__ import annotations

import base64
import html
import io
import json
import re
import tempfile
from datetime import datetime
from typing import Any, Optional
from urllib.parse import urlparse
from pydantic import BaseModel, Field
import markdown
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate


class Action:
    class Valves(BaseModel):
        priority: int = Field(
            default=0,
            description="Controls button display order (lower = appears first)",
        )

        filename_prefix: str = Field(
            default="message",
            description="Prefix used for the downloaded DOCX file name. The final file name is always '<filename_prefix>-<message_id>.docx'.",
        )

        template_url: str = Field(
            default="https://docs.google.com/document/d/1LDmAUhC9wL-qnGS9sZQmlSaF5VSsOcEw7oP1ahM_Al8/edit?usp=sharing",
            description="Publicly accessible URL to the DOCX template file. The server must be able to download it. Google Docs links are supported and automatically converted to DOCX export URLs.",
        )

        body_placeholder: str = Field(
            default="{{ BODY_CONTENT }}",
            description="Placeholder inside the DOCX template where the generated document body will be inserted. This placeholder should appear as a standalone paragraph in the template.",
        )

        mermaid_scale: int = Field(
            default=2,
            description="Rasterization scale used when converting rendered Mermaid diagrams from the browser into PNG images before inserting them into the DOCX.",
        )

        max_mermaid_width_in: float = Field(
            default=6.5,
            description="Maximum width in inches for Mermaid diagrams inserted into the DOCX. Larger diagrams are scaled down proportionally.",
        )

        max_mermaid_height_in: float = Field(
            default=4.5,
            description="Maximum height in inches for Mermaid diagrams inserted into the DOCX. Taller diagrams are scaled down proportionally.",
        )

        max_image_width_in: float = Field(
            default=6.5,
            description="Maximum width in inches for normal images inserted into the DOCX body.",
        )

        max_image_height_in: float = Field(
            default=6.0,
            description="Maximum height in inches for normal images inserted into the DOCX body.",
        )

        request_timeout_s: int = Field(
            default=30,
            description="HTTP timeout in seconds for downloading the template and any externally referenced images.",
        )

    def __init__(self):
        self.valves = self.Valves()

    #
    # ----------------------------
    # Basic helpers
    # ----------------------------
    #

    def build_filename(self, message_id: str) -> str:
        return f"{self.valves.filename_prefix}-{message_id}.docx"

    async def emit_error(self, message: str, __event_emitter__=None):
        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "notification",
                    "data": {"type": "error", "content": message},
                }
            )

    def normalize_template_url(self, url: str) -> str:
        url = (url or "").strip()
        if not url:
            return url

        match = re.match(
            r"^https://docs\.google\.com/document/d/([a-zA-Z0-9_-]+)(?:/.*)?$",
            url,
        )
        if match:
            doc_id = match.group(1)
            return f"https://docs.google.com/document/d/{doc_id}/export?format=docx"

        return url

    def _normalize_content(self, content: Any) -> str:
        if content is None:
            return ""

        if isinstance(content, str):
            return content

        if isinstance(content, list):
            parts: list[str] = []
            for item in content:
                if not isinstance(item, dict):
                    continue
                if item.get("type") == "text":
                    parts.append(item.get("text", ""))
                elif isinstance(item.get("content"), str):
                    parts.append(item.get("content"))
            return "\n".join(part for part in parts if part)

        return str(content)

    def get_message_content(self, body: dict) -> str:
        target_id = body.get("id")
        messages = body.get("messages", []) or []

        if target_id:
            for message in messages:
                if not isinstance(message, dict):
                    continue

                candidate_ids = [
                    message.get("id"),
                    message.get("message_id"),
                    (
                        (message.get("info") or {}).get("id")
                        if isinstance(message.get("info"), dict)
                        else None
                    ),
                    (
                        (message.get("meta") or {}).get("id")
                        if isinstance(message.get("meta"), dict)
                        else None
                    ),
                ]

                if target_id in candidate_ids and message.get("role") == "assistant":
                    return self._normalize_content(message.get("content"))

        for message in reversed(messages):
            if isinstance(message, dict) and message.get("role") == "assistant":
                return self._normalize_content(message.get("content"))

        return ""

    #
    # ----------------------------
    # Placeholder handling
    # ----------------------------
    #

    def get_builtin_placeholder_names(self) -> set[str]:
        return {
            "FILE_NAME",
            "EXPORT_DATE",
            "EXPORT_TIME",
            "USER_NAME",
            "BODY_CONTENT",
        }

    def extract_placeholders_from_template_xml(
        self, template_bytes: bytes
    ) -> list[str]:
        """
        Extract placeholders from the raw .docx XML files.

        We scan the full zip payload textually so placeholders in body, headers,
        footers, tables, etc. can be detected. We intentionally support only
        uppercase snake case placeholders like {{ CLIENT_NAME }}.
        """
        import zipfile

        names: set[str] = set()
        pattern = re.compile(r"\{\{\s*([A-Z][A-Z0-9_]*)\s*\}\}")

        with zipfile.ZipFile(io.BytesIO(template_bytes), "r") as zf:
            for info in zf.infolist():
                name = info.filename
                if not name.endswith(".xml"):
                    continue
                if not (
                    name.startswith("word/")
                    or name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                ):
                    continue

                try:
                    xml_text = zf.read(info).decode("utf-8", errors="ignore")
                except Exception:
                    continue

                for match in pattern.finditer(xml_text):
                    names.add(match.group(1).strip())

        return sorted(names)

    def placeholder_to_label(self, name: str) -> str:
        return name.replace("_", " ").strip().capitalize()

    async def prompt_for_custom_placeholders(
        self,
        placeholder_names: list[str],
        __event_call__=None,
    ) -> dict[str, str]:
        values: dict[str, str] = {}

        if __event_call__ is None:
            for name in placeholder_names:
                values[name] = ""
            return values

        for name in placeholder_names:
            response = await __event_call__(
                {
                    "type": "input",
                    "data": {
                        "title": self.placeholder_to_label(name),
                        "message": f"Enter a value for {name}:",
                        "placeholder": self.placeholder_to_label(name),
                    },
                }
            )

            values[name] = response

        return values

    #
    # ----------------------------
    # Mermaid extraction
    # ----------------------------
    #

    def build_extract_mermaid_png_js(self, message_id: str) -> str:
        scale = self.valves.mermaid_scale
        return f"""
const messageId = {json.dumps(message_id)};
const exportScale = {scale};

function findMessageElement(id) {{
  const selectors = [
    `#message-${{CSS.escape(id)}}`,
    `#${{CSS.escape(id)}}`,
    `[data-message-id="${{id}}"]`,
    `[data-id="${{id}}"]`,
    `.message-${{CSS.escape(id)}}`
  ];
  for (const selector of selectors) {{
    const el = document.querySelector(selector);
    if (el) return el;
  }}
  return null;
}}

function sleep(ms) {{
  return new Promise(resolve => setTimeout(resolve, ms));
}}

async function loadHtml2Canvas() {{
  if (typeof window.html2canvas === "function") return;

  await new Promise((resolve, reject) => {{
    const existing = document.querySelector('script[data-openwebui-mermaid-export="html2canvas-pro"]');
    if (existing) {{
      existing.addEventListener("load", () => resolve(), {{ once: true }});
      existing.addEventListener("error", () => reject(new Error("Failed to load html2canvas-pro")), {{ once: true }});
      return;
    }}

    const script = document.createElement("script");
    script.src = "https://cdn.jsdelivr.net/npm/html2canvas-pro@1.5.12/dist/html2canvas-pro.min.js";
    script.async = true;
    script.dataset.openwebuiMermaidExport = "html2canvas-pro";
    script.onload = () => resolve();
    script.onerror = () => reject(new Error("Failed to load html2canvas-pro"));
    document.head.appendChild(script);
  }});

  if (typeof window.html2canvas !== "function") {{
    throw new Error("html2canvas-pro loaded but window.html2canvas is unavailable.");
  }}
}}

async function waitForImages(root, timeoutMs = 10000) {{
  const images = Array.from(root.querySelectorAll("img"));
  if (!images.length) return;

  await Promise.race([
    Promise.all(images.map(img => {{
      if (img.complete) return Promise.resolve();
      return new Promise(resolve => {{
        const done = () => resolve();
        img.addEventListener("load", done, {{ once: true }});
        img.addEventListener("error", done, {{ once: true }});
      }});
    }})),
    sleep(timeoutMs)
  ]);
}}

const root = findMessageElement(messageId);
if (!root) {{
  throw new Error(`Could not find message element for message id: ${{messageId}}`);
}}

await loadHtml2Canvas();

const diagramElements = Array.from(
  root.querySelectorAll("svg[id^='mermaid-'], svg.flowchart")
);

const diagrams = [];

for (let index = 0; index < diagramElements.length; index++) {{
  const svg = diagramElements[index];

  const tempRoot = document.createElement("div");
  tempRoot.style.position = "fixed";
  tempRoot.style.left = "-100000px";
  tempRoot.style.top = "0";
  tempRoot.style.background = "#ffffff";
  tempRoot.style.padding = "0";
  tempRoot.style.margin = "0";
  tempRoot.style.zIndex = "-1";
  tempRoot.style.display = "inline-block";
  tempRoot.style.boxSizing = "border-box";

  const clone = svg.cloneNode(true);
  clone.removeAttribute("width");
  clone.removeAttribute("height");
  clone.style.display = "block";
  clone.style.margin = "0";
  clone.style.background = "#ffffff";

  const viewBox = svg.getAttribute("viewBox");
  if (viewBox) {{
    const parts = viewBox.trim().split(/\\s+/).map(Number);
    if (parts.length === 4 && Number.isFinite(parts[2]) && Number.isFinite(parts[3])) {{
      clone.setAttribute("width", String(parts[2]));
      clone.setAttribute("height", String(parts[3]));
      tempRoot.style.width = `${{parts[2]}}px`;
      tempRoot.style.height = `${{parts[3]}}px`;
    }}
  }}

  tempRoot.appendChild(clone);
  document.body.appendChild(tempRoot);

  try {{
    await waitForImages(tempRoot, 5000);
    await sleep(150);

    const canvas = await window.html2canvas(tempRoot, {{
      scale: exportScale,
      useCORS: true,
      backgroundColor: "#ffffff",
      logging: false,
      allowTaint: false,
      foreignObjectRendering: false
    }});

    const png = canvas.toDataURL("image/png");

    diagrams.push({{
      index,
      id: svg.id || null,
      png,
      width: canvas.width,
      height: canvas.height
    }});
  }} finally {{
    tempRoot.remove();
  }}
}}

return {{
  success: true,
  messageId,
  diagramCount: diagrams.length,
  diagrams
}};
"""

    def scale_dimensions(
        self,
        width: int,
        height: int,
        max_width: float,
        max_height: float,
    ) -> tuple[float, float]:
        if width <= 0 or height <= 0:
            return max_width, max_height

        width_in = width / 96.0
        height_in = height / 96.0

        ratio = min(
            1.0,
            max_width / width_in if max_width > 0 else 1.0,
            max_height / height_in if max_height > 0 else 1.0,
        )

        return width_in * ratio, height_in * ratio

    def replace_mermaid_blocks_with_markers(
        self,
        markdown_text: str,
        diagrams: list[dict],
    ) -> str:
        pattern = re.compile(
            r"```mermaid[ \t]*\n(.*?)\n```",
            flags=re.IGNORECASE | re.DOTALL,
        )
        diagram_iter = iter(range(len(diagrams)))

        def repl(match: re.Match) -> str:
            try:
                idx = next(diagram_iter)
                return f"\n\n[MERMAID_IMAGE_{idx}]\n\n"
            except StopIteration:
                code = html.escape(match.group(1).strip())
                return f'\n<pre><code class="language-mermaid">{code}</code></pre>\n'

        return pattern.sub(repl, markdown_text)

    #
    # ----------------------------
    # Template and image fetching
    # ----------------------------
    #

    def download_template_bytes(self, template_url: str) -> bytes:
        url = self.normalize_template_url(template_url)
        response = requests.get(url, timeout=self.valves.request_timeout_s)
        response.raise_for_status()
        return response.content

    def data_url_to_bytes(self, data_url: str) -> bytes:
        match = re.match(r"^data:[^;]+;base64,(.+)$", data_url, re.DOTALL)
        if not match:
            raise ValueError("Invalid data URL")
        return base64.b64decode(match.group(1))

    def fetch_image_bytes(self, src: str) -> bytes:
        src = (src or "").strip()
        if not src:
            raise ValueError("Empty image src")

        if src.startswith("data:"):
            return self.data_url_to_bytes(src)

        parsed = urlparse(src)
        if parsed.scheme in {"http", "https"}:
            response = requests.get(src, timeout=self.valves.request_timeout_s)
            response.raise_for_status()
            return response.content

        raise ValueError(f"Unsupported image src: {src}")

    #
    # ----------------------------
    # DOCX body generation
    # ----------------------------
    #

    def add_text_runs(self, paragraph, node, bold=False, italic=False, code=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            return

        if not isinstance(node, Tag):
            return

        tag = node.name.lower()

        if tag == "br":
            paragraph.add_run().add_break()
            return

        if tag in {"strong", "b"}:
            for child in node.children:
                self.add_text_runs(
                    paragraph,
                    child,
                    bold=True,
                    italic=italic,
                    code=code,
                )
            return

        if tag in {"em", "i"}:
            for child in node.children:
                self.add_text_runs(
                    paragraph,
                    child,
                    bold=bold,
                    italic=True,
                    code=code,
                )
            return

        if tag == "code":
            for child in node.children:
                self.add_text_runs(
                    paragraph,
                    child,
                    bold=bold,
                    italic=italic,
                    code=True,
                )
            return

        if tag == "img":
            alt = node.get("alt") or node.get("src") or "[image]"
            run = paragraph.add_run(f"[Image: {alt}]")
            run.italic = True
            return

        for child in node.children:
            self.add_text_runs(
                paragraph,
                child,
                bold=bold,
                italic=italic,
                code=code,
            )

    def insert_paragraph_before(
        self, doc: Document, anchor_paragraph, style: Optional[str] = None
    ):
        p = doc.add_paragraph(style=style)
        anchor_paragraph._element.addprevious(p._element)
        return p

    def insert_image_before(
        self,
        doc: Document,
        anchor_paragraph,
        image_bytes: bytes,
        width_in: float,
        height_in: float,
    ):
        p = doc.add_paragraph()
        run = p.add_run()
        stream = io.BytesIO(image_bytes)

        if width_in > 0 and height_in > 0:
            run.add_picture(stream, width=Inches(width_in), height=Inches(height_in))
        elif width_in > 0:
            run.add_picture(stream, width=Inches(width_in))
        elif height_in > 0:
            run.add_picture(stream, height=Inches(height_in))
        else:
            run.add_picture(stream)

        p.alignment = 1
        anchor_paragraph._element.addprevious(p._element)
        return p

    def insert_table_before(
        self, doc: Document, anchor_paragraph, rows: int, cols: int
    ):
        table = doc.add_table(rows=rows, cols=cols)
        anchor_paragraph._element.addprevious(table._element)
        return table

    def insert_html_image_before(self, doc: Document, anchor_paragraph, src: str):
        image_bytes = self.fetch_image_bytes(src)
        self.insert_image_before(
            doc,
            anchor_paragraph,
            image_bytes=image_bytes,
            width_in=self.valves.max_image_width_in,
            height_in=self.valves.max_image_height_in,
        )

    def insert_heading_before(self, doc: Document, anchor_paragraph, element: Tag):
        level = 1
        try:
            level = int(element.name[1])
        except Exception:
            level = 1
        level = max(1, min(level, 9))

        style_name = f"Heading {level}"
        try:
            p = self.insert_paragraph_before(doc, anchor_paragraph, style=style_name)
        except Exception:
            p = self.insert_paragraph_before(doc, anchor_paragraph)
        self.add_text_runs(p, element)

    def insert_blockquote_before(self, doc: Document, anchor_paragraph, element: Tag):
        style_name = None
        try:
            style_names = [s.name for s in doc.styles]
            if "Intense Quote" in style_names:
                style_name = "Intense Quote"
            elif "Quote" in style_names:
                style_name = "Quote"
        except Exception:
            style_name = None

        p = self.insert_paragraph_before(doc, anchor_paragraph, style=style_name)
        self.add_text_runs(p, element)

    def insert_code_block_before(self, doc: Document, anchor_paragraph, element: Tag):
        p = self.insert_paragraph_before(doc, anchor_paragraph)
        text = element.get_text("\n")
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def insert_list_before(
        self,
        doc: Document,
        anchor_paragraph,
        element: Tag,
        ordered: bool,
        depth: int = 0,
    ):
        style = "List Number" if ordered else "List Bullet"

        for li in element.find_all("li", recursive=False):
            try:
                p = self.insert_paragraph_before(doc, anchor_paragraph, style=style)
            except Exception:
                p = self.insert_paragraph_before(doc, anchor_paragraph)

            if depth > 0:
                p.paragraph_format.left_indent = Inches(0.25 * depth)

            for child in li.contents:
                if isinstance(child, Tag) and child.name.lower() in {"ul", "ol"}:
                    continue
                self.add_text_runs(p, child)

            for nested in li.find_all(["ul", "ol"], recursive=False):
                self.insert_list_before(
                    doc,
                    anchor_paragraph,
                    nested,
                    ordered=(nested.name.lower() == "ol"),
                    depth=depth + 1,
                )

    def insert_html_table_before(self, doc: Document, anchor_paragraph, element: Tag):
        rows = element.find_all("tr")
        if not rows:
            return

        max_cols = 0
        data_rows: list[list[str]] = []
        for tr in rows:
            cells = tr.find_all(["th", "td"], recursive=False)
            row_data = [cell.get_text(" ", strip=True) for cell in cells]
            max_cols = max(max_cols, len(row_data))
            data_rows.append(row_data)

        if max_cols == 0:
            return

        table = self.insert_table_before(
            doc, anchor_paragraph, len(data_rows), max_cols
        )
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for r_idx, row_data in enumerate(data_rows):
            for c_idx in range(max_cols):
                text = row_data[c_idx] if c_idx < len(row_data) else ""
                table.cell(r_idx, c_idx).text = text

    def insert_body_content(
        self,
        doc: Document,
        placeholder: str,
        markdown_text: str,
        diagrams: list[dict],
    ):
        anchor = None
        normalized_placeholder = (placeholder or "").strip()

        for p in doc.paragraphs:
            if normalized_placeholder in (p.text or ""):
                anchor = p
                break

        if anchor is None:
            raise ValueError(
                f"Body placeholder {placeholder!r} was not found in the template."
            )

        html_with_markers = self.replace_mermaid_blocks_with_markers(
            markdown_text, diagrams
        )
        html_body = markdown.markdown(
            html_with_markers,
            extensions=[
                "extra",
                "tables",
                "fenced_code",
                "sane_lists",
                "nl2br",
            ],
            output_format="html5",
        )

        soup = BeautifulSoup(html_body, "html.parser")
        container = soup.body if soup.body else soup

        for element in container.contents:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    p = self.insert_paragraph_before(doc, anchor)
                    p.add_run(text)
                continue

            if not isinstance(element, Tag):
                continue

            tag = element.name.lower()

            if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                self.insert_heading_before(doc, anchor, element)

            elif tag == "p":
                paragraph_text = element.get_text(" ", strip=True)
                mermaid_match = re.fullmatch(
                    r"\[MERMAID_IMAGE_(\d+)\]", paragraph_text or ""
                )

                if mermaid_match:
                    try:
                        idx = int(mermaid_match.group(1))
                        item = diagrams[idx]
                        png = item.get("png", "")
                        width = int(item.get("width", 0) or 0)
                        height = int(item.get("height", 0) or 0)

                        if png.startswith("data:image/png;base64,"):
                            img_bytes = self.data_url_to_bytes(png)
                            scaled_w, scaled_h = self.scale_dimensions(
                                width=width,
                                height=height,
                                max_width=self.valves.max_mermaid_width_in,
                                max_height=self.valves.max_mermaid_height_in,
                            )
                            self.insert_image_before(
                                doc, anchor, img_bytes, scaled_w, scaled_h
                            )
                            continue
                    except Exception:
                        p = self.insert_paragraph_before(doc, anchor)
                        p.add_run("[Mermaid diagram could not be embedded]")
                        continue

                img_tags = element.find_all("img", recursive=False)
                if img_tags and len(list(element.children)) == len(img_tags):
                    for img in img_tags:
                        try:
                            self.insert_html_image_before(
                                doc, anchor, img.get("src", "")
                            )
                        except Exception:
                            p = self.insert_paragraph_before(doc, anchor)
                            p.add_run(
                                f"[Image could not be loaded: {img.get('src', '')}]"
                            )
                    continue

                p = self.insert_paragraph_before(doc, anchor)
                for child in element.children:
                    if isinstance(child, Tag) and child.name.lower() == "img":
                        try:
                            self.insert_html_image_before(
                                doc, anchor, child.get("src", "")
                            )
                        except Exception:
                            p.add_run(
                                f"[Image could not be loaded: {child.get('src', '')}]"
                            )
                    else:
                        self.add_text_runs(p, child)

            elif tag == "blockquote":
                self.insert_blockquote_before(doc, anchor, element)

            elif tag == "pre":
                self.insert_code_block_before(doc, anchor, element)

            elif tag == "ul":
                self.insert_list_before(doc, anchor, element, ordered=False)

            elif tag == "ol":
                self.insert_list_before(doc, anchor, element, ordered=True)

            elif tag == "table":
                self.insert_html_table_before(doc, anchor, element)

            elif tag == "hr":
                p = self.insert_paragraph_before(doc, anchor)
                p.add_run("─" * 24)

            elif tag == "img":
                try:
                    self.insert_html_image_before(doc, anchor, element.get("src", ""))
                except Exception:
                    p = self.insert_paragraph_before(doc, anchor)
                    p.add_run(f"[Image could not be loaded: {element.get('src', '')}]")

            else:
                p = self.insert_paragraph_before(doc, anchor)
                self.add_text_runs(p, element)

        parent = anchor._element.getparent()
        parent.remove(anchor._element)

    #
    # ----------------------------
    # DOCX generation
    # ----------------------------
    #

    def build_context(
        self,
        file_name: str,
        user_name: str,
        custom_placeholders: dict[str, str],
    ) -> dict[str, str]:
        now = datetime.now()
        context = {
            "FILE_NAME": file_name,
            "EXPORT_DATE": now.strftime("%Y-%m-%d"),
            "EXPORT_TIME": now.strftime("%H:%M"),
            "USER_NAME": user_name,
            "BODY_CONTENT": self.valves.body_placeholder,
        }

        for key, value in (custom_placeholders or {}).items():
            context[key] = value or ""

        return context

    def build_docx(
        self,
        template_bytes: bytes,
        markdown_text: str,
        diagrams: list[dict],
        file_name: str,
        user_name: str,
        custom_placeholders: dict[str, str],
    ) -> bytes:
        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=True
        ) as tmp_tpl, tempfile.NamedTemporaryFile(
            suffix=".docx", delete=True
        ) as tmp_rendered:
            tmp_tpl.write(template_bytes)
            tmp_tpl.flush()

            tpl = DocxTemplate(tmp_tpl.name)
            context = self.build_context(
                file_name=file_name,
                user_name=user_name,
                custom_placeholders=custom_placeholders,
            )
            tpl.render(context)
            tpl.save(tmp_rendered.name)

            doc = Document(tmp_rendered.name)
            self.insert_body_content(
                doc=doc,
                placeholder=self.valves.body_placeholder,
                markdown_text=markdown_text,
                diagrams=diagrams,
            )

            out_stream = io.BytesIO()
            doc.save(out_stream)
            return out_stream.getvalue()

    async def download_file(
        self,
        docx_bytes: bytes,
        filename: str,
        __event_emitter__=None,
        __event_call__=None,
    ):
        encoded = base64.b64encode(docx_bytes).decode("ascii")

        js_code = f"""
const base64 = {json.dumps(encoded)};
const filename = {json.dumps(filename)};
const binary = atob(base64);
const bytes = new Uint8Array(binary.length);

for (let i = 0; i < binary.length; i++) {{
  bytes[i] = binary.charCodeAt(i);
}}

const blob = new Blob(
  [bytes],
  {{ type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }}
);
const url = URL.createObjectURL(blob);

try {{
  const a = document.createElement("a");
  a.href = url;
  a.download = filename;
  a.style.display = "none";
  document.body.appendChild(a);
  a.click();
  a.remove();
}} finally {{
  setTimeout(() => URL.revokeObjectURL(url), 4000);
}}

return {{ success: true, filename, size: bytes.length }};
"""

        payload = {"type": "execute", "data": {"code": js_code}}

        if __event_call__ is not None:
            return await __event_call__(payload)

        if __event_emitter__ is not None:
            await __event_emitter__(payload)
            return {"success": True, "filename": filename, "size": len(docx_bytes)}

        return None

    #
    # ----------------------------
    # Main action
    # ----------------------------
    #

    async def action(
        self,
        body: dict,
        __user__=None,
        __event_emitter__=None,
        __event_call__=None,
        **kwargs,
    ):
        message_id = body.get("id")
        if not message_id:
            await self.emit_error(
                "DOCX export failed: could not determine the current message id.",
                __event_emitter__,
            )
            return {
                "content": "Could not determine the current message id from body['id']."
            }

        filename = self.build_filename(message_id)

        template_url = (self.valves.template_url or "").strip()
        if not template_url:
            await self.emit_error(
                "DOCX export failed: template_url is required.",
                __event_emitter__,
            )
            return {"content": "DOCX export failed: template_url is required."}

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": "Reading message content...",
                        "done": False,
                    },
                }
            )

        markdown_text = self.get_message_content(body)
        if not markdown_text.strip():
            await self.emit_error(
                "DOCX export failed: no assistant message content found.",
                __event_emitter__,
            )
            return {"content": "No assistant message content found."}

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": "Downloading DOCX template...",
                        "done": False,
                    },
                }
            )

        try:
            template_bytes = self.download_template_bytes(template_url)
        except Exception as e:
            await self.emit_error(
                f"DOCX export failed: could not download the template file. {e}",
                __event_emitter__,
            )
            return {
                "content": f"DOCX export failed: could not download the template file. {e}"
            }

        all_placeholders = self.extract_placeholders_from_template_xml(template_bytes)
        builtin_placeholders = self.get_builtin_placeholder_names()
        custom_placeholder_names = [
            name for name in all_placeholders if name not in builtin_placeholders
        ]

        custom_placeholders = await self.prompt_for_custom_placeholders(
            custom_placeholder_names,
            __event_call__=__event_call__,
        )

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": "Collecting rendered Mermaid diagrams...",
                        "done": False,
                    },
                }
            )

        diagrams = []
        extract_result = None
        extract_error = None

        try:
            if __event_call__ is None:
                raise RuntimeError(
                    "This action needs __event_call__ so browser JS can return Mermaid image data."
                )

            extract_js = self.build_extract_mermaid_png_js(message_id)
            extract_result = await __event_call__(
                {
                    "type": "execute",
                    "data": {"code": extract_js},
                }
            )

            if not isinstance(extract_result, dict):
                raise RuntimeError(
                    f"Unexpected execute result type: {type(extract_result).__name__}"
                )

            diagrams = extract_result.get("diagrams", []) or []

        except Exception as e:
            extract_error = str(e)
            diagrams = []

        user_name = ""
        if isinstance(__user__, dict):
            user_name = (__user__.get("name") or "").strip()

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "Generating DOCX...", "done": False},
                }
            )

        try:
            docx_bytes = self.build_docx(
                template_bytes=template_bytes,
                markdown_text=markdown_text,
                diagrams=diagrams,
                file_name=filename,
                user_name=user_name,
                custom_placeholders=custom_placeholders,
            )
        except Exception as e:
            await self.emit_error(f"DOCX export failed: {e}", __event_emitter__)
            return {
                "content": f"DOCX export failed: {e}",
                "mermaid_extract_result": extract_result,
                "mermaid_extract_error": extract_error,
            }

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "Starting download...", "done": False},
                }
            )

        result = await self.download_file(
            docx_bytes=docx_bytes,
            filename=filename,
            __event_emitter__=__event_emitter__,
            __event_call__=__event_call__,
        )

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {"description": "DOCX export complete.", "done": True},
                }
            )

        return {
            "content": f"Exported message to DOCX: {filename}",
            "result": result,
            "custom_placeholders": custom_placeholders,
            "mermaid_diagrams_embedded": len(diagrams),
            "mermaid_extract_result": extract_result,
            "mermaid_extract_error": extract_error,
        }